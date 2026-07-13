# -*- coding: utf-8 -*-
"""
TnQ Dashboard - Funnel Quality & Revenue 
Framework: Streamlit
"""

import streamlit as st
import pandas as pd
import requests
import time
import io
import os
import shutil
import zipfile
from datetime import date, timedelta, datetime, timezone

# --- RESILIENT ENVIRONMENT IMPORTS ---
try:
    from google import genai
    HAS_GEMINI = True
except ImportError:
    HAS_GEMINI = False

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# --- 1. CONFIGURATION & CONSTANTS ---
st.set_page_config(page_title="TnQ | Funnel Quality", layout="wide")

REDASH_URL = "https://redash.vahan.link"
QUERY_ID = 17682
ACTIVE_CLIENTS = ["blinkit", "swiggy", "swiggy instamart", "uber"]
CLIENT_FULL = {ck: ck.upper() for ck in ACTIVE_CLIENTS}

# APIs provided by user
REDASH_API_KEY = "4aFm2iOoyx8I91svQccdeZr0jmaiUsMFSRinZcmu"
GEMINI_API_KEY = "4aFm2iOoyx8I91svQccdeZr0jmaiUsMFSRinZcmu"

CLIENT_MS = {
    "blinkit": ["20th", "60th", "100th", "120th", "150th", "200th"],
    "swiggy": ["5th", "10th", "20th", "50th", "60th", "80th", "100th", "150th", "200th"],
    "swiggy instamart": ["5th", "10th", "20th", "50th", "60th", "80th", "100th", "150th", "200th"],
    "uber": ["10th", "20th", "30th", "50th", "100th", "150th", "200th"],
}

CLIENT_KEY_MS = {
    "blinkit": "60th", "swiggy": "20th", "swiggy instamart": "20th", "uber": "20th",
}
MIN_VL_FODS = 0

# --- FINANCIAL & RISK THRESHOLDS ---
MIN_CURRENT_MTD_FODS = 25
DROP_CRITICAL = -15
DROP_HIGH = -8
DROP_WATCH = -3
ABS_CRITICAL = 0.30
ABS_HIGH = 0.50
ABS_WATCH = 0.70
LT_CRITICAL = 5
LT_HIGH = 10
SURGE_THRESH = 100
SURGE_DROP = -5
BELOW20_WATCH = 65

CLIENT_DECLINE_MS = {
    "blinkit":          ("20th", "60th"),
    "swiggy":           ("20th", "50th"),
    "swiggy instamart": ("20th", "50th"),
    "uber":             ("10th", "20th"),
}

MISUSE_SHOW_MS = {
    "swiggy":           ["5th", "20th", "50th", "100th"],
    "swiggy instamart": ["20th", "50th", "100th"],
    "blinkit":          ["20th", "60th", "100th", "200th"],
    "uber":             ["10th", "20th", "30th", "50th"],
}

# --- FINANCIAL RATE CARDS & CONSTANTS ---
BLINKIT_CRITICAL_CITIES = {
    "delhi", "mumbai", "bangalore", "hyderabad", "pune", "kolkata", "chennai",
    "ahmedabad", "jaipur", "lucknow", "gurgaon", "noida", "indore", "chandigarh",
    "ghaziabad", "faridabad"
}

COMMERCIALS = {
    "swiggy":               {"20th": 400, "50th": 700, "60th": 0,    "80th": 900,  "100th": 1000, "120th": 0,    "150th": 0, "200th": 0},
    "swiggy instamart":     {"20th": 400, "50th": 750, "60th": 0,    "80th": 1000, "100th": 1100, "120th": 0,    "150th": 0, "200th": 0},
    "blinkit_critical":     {"20th": 860, "50th": 0,   "60th": 1130, "80th": 0,    "100th": 1200, "120th": 1560, "150th": 0, "200th": 4750},
    "blinkit_non_critical": {"20th": 690, "50th": 0,   "60th": 780,  "80th": 0,    "100th": 800,  "120th": 1040, "150th": 0, "200th": 3310}
}

OFFER_1 = {
    "swiggy":               {"20th": 800,  "50th": 1100, "60th": 0,   "120th": 0},
    "swiggy instamart":     {"20th": 1000, "50th": 1200, "60th": 0,   "120th": 0},
    "blinkit_critical":     {"20th": 1300, "50th": 0,    "60th": 500, "120th": 500},
    "blinkit_non_critical": {"20th": 1150, "50th": 0,    "60th": 400, "120th": 400}
}

TARGET_DIP_MS = ["20th", "50th", "60th", "80th", "100th", "120th", "150th", "200th"]

def get_segment(client, city=None):
    if client == "blinkit":
        city_clean = str(city).strip().lower() if pd.notna(city) else ""
        return "blinkit_critical" if city_clean in BLINKIT_CRITICAL_CITIES else "blinkit_non_critical"
    return client

def fmt_currency(val):
    if pd.isna(val): return "₹0"
    if val < 0: return f"-₹{abs(val):,.0f}"
    return f"₹{val:,.0f}"

# Date Calculations
yesterday = date.today() - timedelta(days=6)
mtd_day = yesterday.day
start_month = yesterday.month - 3
start_year = yesterday.year
if start_month <= 0:
    start_month += 12
    start_year -= 1
START_DATE = str(date(start_year, start_month, 1))
END_DATE = str(yesterday)

# --- 2. DATA ACQUISITION & CACHING ---
IST = timezone(timedelta(hours=5, minutes=30))

def get_daily_refresh_key():
    """Generates a unique cache key that updates exactly at 13:30 (1:30 PM) IST every day."""
    now = datetime.now(IST)
    if now.hour < 13 or (now.hour == 13 and now.minute < 30):
        return str(now.date() - timedelta(days=1))
    return str(now.date())

@st.cache_data(show_spinner=False)
def fetch_redash(refresh_key):
    body_fresh = {"parameters": {"Client": ACTIVE_CLIENTS}, "max_age": 7200}
    body_cached = {**body_fresh, "max_age": 7200}
    
    r = requests.post(f"{REDASH_URL}/api/queries/{QUERY_ID}/results?api_key={REDASH_API_KEY}", json=body_fresh, timeout=30)
    j = r.json()
    
    if "query_result" in j:
        return j["query_result"]["data"]["rows"]
    
    if "job" not in j:
        st.error(f"Redash API Error: {j}")
        return []
        
    job_id = j["job"]["id"]
    for attempt in range(40):
        time.sleep(15)
        r2 = requests.post(f"{REDASH_URL}/api/queries/{QUERY_ID}/results?api_key={REDASH_API_KEY}", json=body_cached, timeout=30)
        j2 = r2.json()
        if "query_result" in j2:
            return j2["query_result"]["data"]["rows"]
            
    st.error("Timed out waiting for Redash.")
    return []

# --- 3. DATA PROCESSING ---
@st.cache_data(show_spinner=False)
def run_analysis(rows):
    if not rows: return {}, pd.DataFrame()
    
    df = pd.DataFrame(rows)
    df["_fod"] = pd.to_datetime(df["first_date_of_work"], format="%Y-%m-%d", errors="coerce")
    valid = df["_fod"].notna() & (df["_fod"].dt.day <= mtd_day) & (df["_fod"] <= pd.Timestamp(END_DATE))
    df = df[valid].copy()
    df["_month"] = df["_fod"].dt.strftime("%b-%Y")
    df = df.drop_duplicates(subset=["phone_number", "_month"])
    df["_vl"] = df["vl_name"].fillna("Unknown")
    
    col_map = {str(c).strip().lower(): c for c in df.columns}
    df["ZM"] = df[col_map["zm"]].fillna("Unknown") if "zm" in col_map else "Unknown"
    df["Region"] = df[col_map["region"]].fillna("Unknown") if "region" in col_map else "Unknown"
    df["CL"] = df[col_map["cl"]].fillna("Unknown") if "cl" in col_map else "Unknown"
    
    if "cm" in col_map:
        df["CM"] = df[col_map["cm"]].fillna("Unknown")
    elif "am" in col_map:
        df["CM"] = df[col_map["am"]].fillna("Unknown")
    else:
        df["CM"] = "Unknown"

    for ms in TARGET_DIP_MS:
        col = f"{ms}_order_date"
        if col in df.columns:
            df[col + "_dt"] = pd.to_datetime(df[col], format="%Y-%m-%d", errors="coerce")
            df[f"has_{ms}"] = (
                (df[col + "_dt"].dt.year == df["_fod"].dt.year) &
                (df[col + "_dt"].dt.month == df["_fod"].dt.month) &
                (df[col + "_dt"].dt.day <= mtd_day)
            ).astype(int)

    results = {}
    for client in ACTIVE_CLIENTS:
        sub = df[df["company_name"].str.lower() == client].copy()
        ms_list = CLIENT_MS.get(client, [])
        key_ms = CLIENT_KEY_MS.get(client, ms_list[0])

        for ms in ms_list:
            col = f"{ms}_order_date"
            if col not in sub.columns: sub[col] = None
            sub[col + "_dt"] = pd.to_datetime(sub[col], format="%Y-%m-%d", errors="coerce")
            sub[f"has_{ms}"] = (
                (sub[col + "_dt"].dt.year == sub["_fod"].dt.year) &
                (sub[col + "_dt"].dt.month == sub["_fod"].dt.month) &
                (sub[col + "_dt"].dt.day <= mtd_day)
            ).astype(int)

        all_months = sorted(sub["_month"].unique(), key=lambda x: pd.to_datetime("01 " + x))

        monthly = []
        for m in all_months:
            g = sub[sub["_month"] == m]
            if len(g) == 0: continue
            lt = g["candidate_lifetime_orders_trips"].astype(float)
            rec = {"month": m, "fods": len(g)}
            for ms in ms_list:
                rec[f"pct_{ms}"] = round(g[f"has_{ms}"].mean() * 100, 2)
            rec["avg_lt"] = round(lt.mean(), 2)
            rec["median_lt"] = round(lt.median(), 2)
            rec["pct_200plus"] = round((lt >= 200).mean() * 100, 2)
            rec["pct_below20"] = round((lt < 20).mean() * 100, 2)
            monthly.append(rec)

        bm_row = {
            "VL Name": "⬛ BENCHMARK (MTD)",
            "ZM": "", "Region": "", "CM": "", "CL": "",
            "Total FODs": sum(m["fods"] for m in monthly)
        }
        for ms2 in ms_list:
            vals = [m.get(f"pct_{ms2}", 0) for m in monthly]
            bm_row[f"F{ms2}%"] = round(sum(vals)/len(vals), 2) if vals else 0
        for f2, lbl in [("avg_lt", "Avg LT"), ("median_lt", "Median LT"), ("pct_200plus", "% 200+ LT"), ("pct_below20", "% <20 LT")]:
            vals = [m.get(f2, 0) for m in monthly if m.get(f2) is not None]
            bm_row[lbl] = round(sum(vals)/len(vals), 2) if vals else 0

        bm_ms = {ms2: round(sum(m.get(f"pct_{ms2}", 0) for m in monthly) / max(len(monthly), 1), 2) for ms2 in ms_list}

        vl_summary = []
        vl_monthly = {}
        
        for vl_name, vl_df in sub.groupby("_vl"):
            if len(vl_df) < MIN_VL_FODS: continue
            
            zm_val = vl_df["ZM"].mode()[0] if not vl_df["ZM"].empty else "Unknown"
            reg_val = vl_df["Region"].mode()[0] if not vl_df["Region"].empty else "Unknown"
            cm_val = vl_df["CM"].mode()[0] if not vl_df["CM"].empty else "Unknown"
            cl_val = vl_df["CL"].mode()[0] if not vl_df["CL"].empty else "Unknown"
            
            lt_all = vl_df["candidate_lifetime_orders_trips"].astype(float)
            rec = {
                "vl": vl_name,
                "ZM": zm_val,
                "Region": reg_val,
                "CM": cm_val,
                "CL": cl_val,
                "total_fods": len(vl_df),
                "avg_lt": round(lt_all.mean(), 2),
                "median_lt": round(lt_all.median(), 2),
                "pct_200plus": round((lt_all >= 200).mean() * 100, 2),
                "pct_below20": round((lt_all < 20).mean() * 100, 2),
            }
            for ms in ms_list:
                rec[f"pct_{ms}"] = round(vl_df[f"has_{ms}"].mean() * 100, 2)
                
            vm = {}
            for m in all_months:
                m_df = vl_df[vl_df["_month"] == m]
                if len(m_df) < 5: 
                    vm[m] = None
                    continue
                lt_m = m_df["candidate_lifetime_orders_trips"].astype(float)
                m_rec = {"fods": len(m_df)}
                for ms in ms_list:
                    m_rec[f"pct_{ms}"] = round(m_df[f"has_{ms}"].mean() * 100, 2)
                m_rec["median_lt"] = round(lt_m.median(), 2)
                m_rec["pct_200plus"] = round((lt_m >= 200).mean() * 100, 2)
                m_rec["pct_below20"] = round((lt_m < 20).mean() * 100, 2)
                vm[m] = m_rec
            
            vl_monthly[vl_name] = vm
            
            valid_months = [m for m in all_months if vm.get(m) is not None]
            if len(valid_months) >= 2:
                pm, cm = valid_months[-2], valid_months[-1]
                rec["fod_growth"] = round((vm[cm]["fods"] - vm[pm]["fods"]) / max(vm[pm]["fods"], 1) * 100, 2)
                for ms in ms_list:
                    rec[f"delta_{ms}"] = round(vm[cm].get(f"pct_{ms}", 0) - vm[pm].get(f"pct_{ms}", 0), 2)
                    
            vl_summary.append(rec)

        results[client] = {
            "monthly": monthly,
            "vl_summary": vl_summary,
            "vl_monthly": vl_monthly,
            "bm_ms": bm_ms,
            "bm_row": bm_row,
            "milestones": ms_list,
            "key_ms": key_ms,
        }
    return results, df

@st.cache_data(show_spinner=False)
def calculate_financials(df_raw, results_dict):
    fin_data = {}
    for ck in ["swiggy", "swiggy instamart", "blinkit"]:
        if ck not in results_dict: continue
        
        mon = results_dict[ck]["monthly"]
        if len(mon) < 2: continue
        curr_m, prev_m = mon[-1]["month"], mon[-2]["month"]
        
        if ck in ["swiggy", "swiggy instamart"]:
            ms_list = [m for m in ["20th", "50th", "80th", "100th"] if m in TARGET_DIP_MS]
        else:
            ms_list = [m for m in ["20th", "60th", "120th", "200th"] if m in TARGET_DIP_MS]
            
        group_cols = ["company_name", "_vl", "jobCity"] if ck == "blinkit" else ["company_name", "_vl"]
        df_client = df_raw[(df_raw["company_name"].str.lower() == ck) & (df_raw["_month"].isin([curr_m, prev_m]))]
        
        rows_list = []
        for name, group in df_client.groupby(group_cols):
            vln = name[1]
            city = name[2] if ck == "blinkit" else None
            segment = get_segment(ck, city)
            
            region = group["Region"].mode()[0] if not group["Region"].empty else "Unknown"
            zm = group["ZM"].mode()[0] if not group["ZM"].empty else "Unknown"
            
            c_data = group[group["_month"] == curr_m]
            p_data = group[group["_month"] == prev_m]
            c_fods, p_fods = len(c_data), len(p_data)
            if c_fods == 0 and p_fods == 0: continue
            
            row = {"Client": ck.title(), "VL Name": vln, "Region": region, "ZM": zm, "City": city, "Segment": segment, f"FODs {prev_m[:3]}": p_fods, f"FODs {curr_m[:3]}": c_fods}
            
            for m2 in ms_list:
                col_has = f"has_{m2}"
                c_hits = c_data[col_has].sum() if c_fods > 0 and col_has in c_data.columns else 0
                p_hits = p_data[col_has].sum() if p_fods > 0 and col_has in p_data.columns else 0
                
                rate_rc1 = COMMERCIALS.get(segment, {}).get(m2, 0)
                rate_o1 = OFFER_1.get(segment, {}).get(m2, 0)
                
                row[f"F{m2} Hits {curr_m[:3]}"] = c_hits
                row[f"F{m2} RC1 Rev {curr_m[:3]}"] = fmt_currency(c_hits * rate_rc1)
                row[f"F{m2} RC1 Rev Loss"] = fmt_currency((c_hits * rate_rc1) - (p_hits * rate_rc1))
                
                row[f"F{m2} Offer1 Rev {curr_m[:3]}"] = fmt_currency(c_hits * rate_o1)
                row[f"F{m2} Offer1 Rev Loss"] = fmt_currency((c_hits * rate_o1) - (p_hits * rate_o1))

            rows_list.append(row)
        
        if rows_list:
            fin_data[ck] = pd.DataFrame(rows_list)
            
    return fin_data

# --- WORD DOCUMENT GENERATION ---
def set_thick_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    tblBorders = tblPr.first_child_found_in("w:tblBorders")
    if tblBorders is None:
        tblBorders = OxmlElement('w:tblBorders')
        tblPr.append(tblBorders)
    else:
        tblBorders.clear()

    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '12')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)

def _fmt_pct_word(val):
    return "-" if pd.isna(val) or val is None else f"{val:.1f}%"

def get_30d_misuse_data(df_raw, client, ms_list, end_date_str):
    end_dt = pd.to_datetime(end_date_str)
    start_dt = end_dt - pd.Timedelta(days=30)
    df_30d = df_raw[(df_raw["company_name"].str.lower() == client) & 
                    (df_raw["_fod"] >= start_dt) & 
                    (df_raw["_fod"] <= end_dt)]
                    
    bm_30d = {}
    for m2 in ms_list:
        col_has = f"has_{m2}"
        if col_has in df_30d.columns and len(df_30d) > 0:
            bm_30d[m2] = df_30d[col_has].mean() * 100
        else:
            bm_30d[m2] = 0
            
    critical_vls = []
    for vln, grp in df_30d.groupby("_vl"):
        if len(grp) <= MIN_CURRENT_MTD_FODS: continue
        
        zm = grp["ZM"].mode()[0] if not grp["ZM"].empty else "Unknown"
        lt_all = grp["candidate_lifetime_orders_trips"].astype(float)
        med_lt = lt_all.median()
        
        vl_rec = {
            "vl": vln,
            "zm": zm,
            "fods": len(grp),
            "median_lt": med_lt,
            "is_critical": False
        }
        
        is_critical = False
        for m2 in ms_list:
            col_has = f"has_{m2}"
            vl_pct = grp[col_has].mean() * 100 if col_has in grp.columns else 0
            vl_rec[f"pct_{m2}"] = vl_pct
            vl_rec[f"base_{m2}"] = bm_30d[m2]
            
            if bm_30d[m2] > 0 and (bm_30d[m2] - vl_pct) / bm_30d[m2] >= 0.50:
                is_critical = True
                
        if med_lt < LT_CRITICAL:
            is_critical = True
            
        vl_rec["is_critical"] = is_critical
        if is_critical:
            critical_vls.append(vl_rec)
            
    return {"bm": bm_30d, "vls": critical_vls}

def generate_zm_email_drafts(results, df_raw, mtd_day_val, end_date_str):
    if not HAS_DOCX:
        raise Exception("python-docx is not installed.")

    output_dir = f"ZM_Drafts_{end_date_str}"
    os.makedirs(output_dir, exist_ok=True)
    
    unique_zms = set()
    for ck in ACTIVE_CLIENTS:
        if ck in results:
            for vl in results[ck]["vl_summary"]:
                if vl.get("ZM") and vl["ZM"] != "Unknown":
                    unique_zms.add(str(vl["ZM"]).strip())
    
    unique_zms = sorted(list(unique_zms))
    cohort_month = pd.to_datetime(end_date_str).strftime('%B')
    cohort_day = pd.to_datetime(end_date_str).day

    for zm_name in unique_zms:
        doc = Document()
        p = doc.add_paragraph()
        p.add_run(f"Hi {zm_name},").bold = True
        doc.add_paragraph(f"Please find {cohort_month}'s TnQ quality Report for your cluster at the client level below. Please work with the VLs listed below to improve quality, and share your action plans and the estimated timeframe for improvement.")
        
        note_p = doc.add_paragraph()
        note_run = note_p.add_run(f"Note- 5 additional days have been added for making the report accurate, i.e. FT<= {cohort_month} {cohort_day} but F20 and other milestone have been given {mtd_day_val} addition buffer days i.e. {cohort_month} {cohort_day}+{mtd_day_val} days).")
        note_run.italic = True

        has_content = False

        for ck in CLIENT_ORDER:
            if ck not in results: continue
            data = results[ck]
            client_label = CLIENT_FULL.get(ck, ck.upper())

            # TABLE 1 (MTD VS LMD)
            mon = data["monthly"]
            if len(mon) < 2: continue
            curr_m, prev_m = mon[-1]["month"], mon[-2]["month"]
            ms1, ms2 = CLIENT_DECLINE_MS.get(ck, (data["milestones"][0], data["key_ms"]))

            t1_rows = []
            for vl in data["vl_summary"]:
                if vl.get("ZM") != zm_name: continue
                vln = vl["vl"]
                vm = data["vl_monthly"].get(vln, {})
                curr_d, prev_d = vm.get(curr_m) or {}, vm.get(prev_m) or {}

                curr_f1, prev_f1 = curr_d.get(f"pct_{ms1}"), prev_d.get(f"pct_{ms1}")
                curr_f2, prev_f2 = curr_d.get(f"pct_{ms2}"), prev_d.get(f"pct_{ms2}")
                d_f1 = round(curr_f1 - prev_f1, 1) if curr_f1 is not None and prev_f1 is not None else None
                d_f2 = round(curr_f2 - prev_f2, 1) if curr_f2 is not None and prev_f2 is not None else None

                if (d_f1 is not None and d_f1 < 0) or (d_f2 is not None and d_f2 < 0):
                    t1_rows.append([
                        str(vln), str(zm_name), f"{curr_d.get('fods', 0):,}", f"{prev_d.get('fods', 0):,}",
                        _fmt_pct_word(curr_f1), _fmt_pct_word(prev_f1), _fmt_pct_word(curr_f2), _fmt_pct_word(prev_f2),
                        f"{d_f1:+.1f}%" if d_f1 is not None else "-", f"{d_f2:+.1f}%" if d_f2 is not None else "-"
                    ])

            # TABLE 2 (Platform Avg vs VL Performance using Last 30 Days Logic)
            t2_ms_list = ["20th", "60th", "100th", "200th"]
            misuse_data = get_30d_misuse_data(df_raw, ck, t2_ms_list, end_date_str)
            critical_vls = [v for v in misuse_data['vls'] if v['zm'] == zm_name]

            t2_rows = []
            for vl_rec in critical_vls:
                row_data = [
                    str(vl_rec['vl']), str(vl_rec['zm']), "❌ CRITICAL",
                    f"{vl_rec['fods']:,}", str(round(vl_rec['median_lt'], 1))
                ]
                for m2 in t2_ms_list:
                    if m2 in data["milestones"]:
                        row_data.append(_fmt_pct_word(vl_rec[f"pct_{m2}"]))
                        row_data.append(_fmt_pct_word(vl_rec[f"base_{m2}"]))
                    else:
                        row_data.extend(["-", "-"])
                t2_rows.append(row_data)

            if not t1_rows and not t2_rows: continue
            has_content = True

            client_h = doc.add_heading(client_label, level=2)
            client_h.runs[0].font.color.rgb = RGBColor(197, 90, 0)

            if t1_rows:
                doc.add_heading("MTD VS LMD report", level=3)
                t1_headers = [
                    "VL Name", "ZM Name", f"{curr_m[:3]} MTD FOD", f"LMTD FOD",
                    f"{curr_m[:3]} MTD F{ms1}%", f"LMTD F{ms1}%", f"{curr_m[:3]} F{ms2}%", f"LMTD F{ms2}%",
                    f"Delta F{ms1}", f"Delta F{ms2}"
                ]
                table = doc.add_table(rows=1, cols=len(t1_headers))
                set_thick_borders(table)
                hdr_cells = table.rows[0].cells
                for i, h in enumerate(t1_headers):
                    hdr_cells[i].text = h
                    hdr_cells[i].paragraphs[0].runs[0].font.bold = True
                for row_data in t1_rows:
                    row_cells = table.add_row().cells
                    for i, val in enumerate(row_data):
                        row_cells[i].text = str(val)
                        if i >= 2: row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                doc.add_paragraph()

            if t2_rows:
                doc.add_heading("Platform Avg(Baseline) vs VL Performance report (Last 30 Days)", level=3)
                t2_note = doc.add_paragraph("Note: This table shows the list of VLs whose milestones achieved are critically below the platform average.")
                t2_note.runs[0].italic = True

                t2_headers = [
                    "VL Name", "ZM", "Severity", "Total FODs", "Median LT",
                    "F20th%\n(30d Overall)", "F20th%\n(30d Baseline)", "F60th%\n(30d Overall)", "F60th%\n(30d Baseline)",
                    "F100th%\n(30d Overall)", "F100th%\n(30d Baseline)", "F200th%\n(30d Overall)", "F200th%\n(30d Baseline)"
                ]
                table = doc.add_table(rows=1, cols=len(t2_headers))
                set_thick_borders(table)
                hdr_cells = table.rows[0].cells
                for i, h in enumerate(t2_headers):
                    hdr_cells[i].text = h
                    hdr_cells[i].paragraphs[0].runs[0].font.bold = True
                for row_data in t2_rows:
                    row_cells = table.add_row().cells
                    for i, val in enumerate(row_data):
                        row_cells[i].text = str(val)
                        if i >= 3: row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                doc.add_paragraph()

        if not has_content:
            doc.add_paragraph("No critical flags or negative quality decline metrics for your cluster this month.")

        safe_zm_name = "".join([c for c in zm_name if c.isalpha() or c.isdigit() or c==' ']).rstrip()
        file_path = os.path.join(output_dir, f"Draft_{safe_zm_name.replace(' ', '_')}.docx")
        doc.save(file_path)

    # Zip the generated directory in memory
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zipf:
        for root, _, files in os.walk(output_dir):
            for file in files:
                zipf.write(os.path.join(root, file), file)
    
    shutil.rmtree(output_dir)
    zip_buffer.seek(0)
    return zip_buffer


def draft_summary(results):
    if not HAS_GEMINI:
        return (
            "⚠️ **AI Insights Configuration Missing:**\n"
            "The dependency module `google-genai` was not detected in this Python execution environment.\n\n"
            "**To fix this on Streamlit Cloud:**\n"
            "1. Please add `google-genai` inside your repository's `requirements.txt` file.\n"
            "2. Streamlit will detect the change, automatically rebuild, and activate this panel."
        )
    try:
        client = genai.Client(api_key=GEMINI_API_KEY)
        prompt = "Write a short analytical executive summary for a dashboard about gig worker funnel quality. Highlight main drops in conversion, notable surges, and top client observations. Use 3-4 professional bullet points. No fluff."
        response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=prompt
        )
        return response.text
    except Exception as e:
        return f"Gemini API Error: {e}"

# --- STYLER FUNCTIONS (PANDAS FORMATTING) ---
def highlight_summary(row):
    styles = [''] * len(row)
    if row.name % 2 == 0:
        styles = ['background-color: #F2F2F2; color: #000000'] * len(row)
    if 'Change' in row.index:
        idx = row.index.get_loc('Change')
        val = str(row['Change'])
        if val.startswith('-'):
            styles[idx] += '; color: #C00000; font-weight: bold'
        elif val.startswith('+'):
            styles[idx] += '; color: #375623; font-weight: bold'
    return styles

def highlight_vl_summary(row, bm_dict):
    styles = [''] * len(row)
    if "BENCHMARK" in str(row.get("VL Name", "")):
        return ['background-color: #2C2C2C; color: #FFFFFF; font-weight: bold'] * len(row)
        
    for i, col in enumerate(row.index):
        if isinstance(col, str) and col.startswith('F') and col.endswith('%'):
            val = row[col]
            ms_key = col[1:-1]
            bv = bm_dict.get(ms_key, 0)
            if pd.notna(val) and bv > 0:
                ratio = float(val) / bv
                if ratio >= 1.15: styles[i] = 'background-color: #CCFFCC; color: #375623'
                elif ratio < 0.50: styles[i] = 'background-color: #FFCCCC; color: #C00000'
                elif ratio < 0.80: styles[i] = 'background-color: #FFE4CC; color: #C55A00'
    return styles

def highlight_deltas(row):
    styles = [''] * len(row)
    for i, col in enumerate(row.index):
        if isinstance(col, str) and (col.startswith('Δ') or col.startswith('Delta') or col == 'FOD Growth %' or 'Δ' in col):
            val = row[col]
            if pd.notna(val):
                if val <= -15: styles[i] = 'background-color: #FFCCCC; color: #C00000'
                elif val <= -5: styles[i] = 'background-color: #FFE4CC; color: #C55A00'
                elif val >= 0: styles[i] = 'background-color: #CCFFCC; color: #375623'
    return styles

def highlight_severity_rows(row):
    styles = [''] * len(row)
    if "BENCHMARK" in str(row.get("VL Name", "")):
        return ['background-color: #2C2C2C; color: #FFFFFF; font-weight: bold'] * len(row)
        
    sev = str(row.get("Severity", ""))
    if "CRITICAL" in sev:
        return ['background-color: #FFD2D2; color: #8B0000'] * len(row)
    elif "HIGH" in sev:
        return ['background-color: #FFEED2; color: #8B5A00'] * len(row)
    elif "WATCH" in sev:
        return ['background-color: #FFFFD2; color: #8B8B00'] * len(row)
    return styles

def highlight_misuse_status(val):
    if isinstance(val, str) and "⚠️ DROP" in val:
        return 'color: #C00000; font-weight: bold'
    elif isinstance(val, str) and "✓ OK" in val:
        return 'color: #375623'
    return ''

def style_financials(val):
    if isinstance(val, str) and '-₹' in val:
        return 'color: #C00000; font-weight: bold'
    return ''

# --- 4. STREAMLIT UI (MECE FRAMEWORK) ---
CLIENT_ORDER = ACTIVE_CLIENTS

def main():
    st.title("📊 TnQ: Funnel Quality Hub")
    st.markdown(f"**Data Period:** {START_DATE} → {END_DATE} | **MTD Cutoff:** Day {mtd_day}")
    st.info("🕒 **Data Refresh Schedule:** This dashboard synchronizes with the scheduled Redash pipeline and updates exactly at **1:30 PM IST** daily.")

    with st.spinner("Fetching and processing data pipelines..."):
        refresh_key = get_daily_refresh_key()
        rows = fetch_redash(refresh_key)
        results, df_raw = run_analysis(rows)
        fin_data = calculate_financials(df_raw, results)

    if not results:
        st.warning("No data returned from queries.")
        return

    # --- TAB GENERATION ---
    tab_names = [c.title() for c in ACTIVE_CLIENTS] + ["💰 Commercials"]
    tabs = st.tabs(tab_names)

    for idx, client in enumerate(ACTIVE_CLIENTS):
        with tabs[idx]:
            client_data = results.get(client, {})
            if not client_data:
                st.info("No active data for this client in the current timeframe.")
                continue
                
            ms_list = client_data["milestones"]
            bm_ms = client_data.get("bm_ms", {})
            key_ms = client_data["key_ms"]
            vl_monthly = client_data.get("vl_monthly", {})
            
            df_vl = pd.DataFrame(client_data["vl_summary"])

            # --- ZM & REGION FILTERS ---
            st.markdown(f"### 🔍 Filter Data for {client.title()}")
            col1, col2 = st.columns(2)
            
            regions = ["All"] + sorted([str(x) for x in df_vl["Region"].unique()])
            sel_reg = col1.selectbox("Filter by Region", regions, key=f"reg_{client}")
            if sel_reg != "All":
                df_vl = df_vl[df_vl["Region"] == sel_reg]
                
            zms = ["All"] + sorted([str(x) for x in df_vl["ZM"].unique()])
            sel_zm = col2.selectbox("Filter by ZM", zms, key=f"zm_{client}")
            if sel_zm != "All":
                df_vl = df_vl[df_vl["ZM"] == sel_zm]

            df_vl = df_vl.sort_values(by="total_fods", ascending=False)
            filtered_vl_names = df_vl["vl"].tolist()

            # --- EXPANDER 1: Overall Monthly (Transposed Layout) ---
            with st.expander("📈 Overall Funnel - Month over Month (Unfiltered)", expanded=False):
                if client_data["monthly"]:
                    m_data = client_data["monthly"]
                    all_mths = [m["month"] for m in m_data]
                    
                    metric_rows = [("Total FODs", "fods", "int")]
                    for ms in ms_list: metric_rows.append((f"F{ms}%", f"pct_{ms}", "pct"))
                    metric_rows += [
                        ("Avg LT", "avg_lt", "float"),
                        ("Median LT", "median_lt", "float"),
                        ("% 200+ LT", "pct_200plus", "pct"),
                        ("% <20 LT", "pct_below20", "pct")
                    ]
                    
                    summary_table = []
                    for lbl, field, dtype in metric_rows:
                        row_dict = {"Metric": lbl}
                        for m_dict in m_data:
                            row_dict[m_dict["month"]] = m_dict.get(field)
                        
                        if len(all_mths) >= 2:
                            val1 = m_data[-2].get(field)
                            val2 = m_data[-1].get(field)
                            if val1 is not None and val2 is not None:
                                row_dict["Change"] = val2 - val1
                            else:
                                row_dict["Change"] = None
                        row_dict["_dtype"] = dtype
                        summary_table.append(row_dict)
                    
                    df_summ = pd.DataFrame(summary_table)
                    
                    def format_metric(row):
                        dtype = row["_dtype"]
                        fmt_row = row.copy()
                        for c in df_summ.columns:
                            if c not in ["Metric", "_dtype", "Change"]:
                                val = row[c]
                                if pd.isna(val): fmt_row[c] = "-"
                                elif dtype == "int": fmt_row[c] = f"{int(val):,}"
                                elif dtype == "pct": fmt_row[c] = f"{val:.2f}%"
                                elif dtype == "float": fmt_row[c] = f"{val:.2f}"
                            if c == "Change":
                                val = row[c]
                                if pd.isna(val): fmt_row[c] = "-"
                                elif dtype == "int": fmt_row[c] = f"{int(val):+,}"
                                elif dtype in ["pct", "float"]: fmt_row[c] = f"{val:+.2f}"
                                if dtype == "pct" and not pd.isna(val): fmt_row[c] += " pp"
                        return fmt_row
                    
                    df_summ_fmt = df_summ.apply(format_metric, axis=1).drop(columns=["_dtype"])
                    st.dataframe(df_summ_fmt.style.apply(highlight_summary, axis=1), width="stretch", hide_index=True)

            # --- EXPANDER 2: VL Summary (Includes Benchmark Row) ---
            with st.expander("🏢 VL Summary (Current MTD vs Benchmark)", expanded=True):
                ms_cols = [f"pct_{m}" for m in ms_list]
                disp_cols1 = ["vl", "ZM", "Region", "CM", "CL", "total_fods", "avg_lt", "median_lt", "pct_200plus", "pct_below20"] + ms_cols
                disp_cols1 = [c for c in disp_cols1 if c in df_vl.columns]
                
                df_disp1 = df_vl[disp_cols1].copy()
                rename_map1 = {"vl": "VL Name", "total_fods": "Total FODs", "avg_lt": "Avg LT", "median_lt": "Median LT", "pct_200plus": "% 200+ LT", "pct_below20": "% <20 LT"}
                rename_map1.update({f"pct_{m}": f"F{m}%" for m in ms_list})
                df_disp1.rename(columns=rename_map1, inplace=True)
                
                bm_df = pd.DataFrame([client_data.get("bm_row", {})])
                df_disp1 = pd.concat([bm_df, df_disp1], ignore_index=True)
                
                st.dataframe(df_disp1.style.apply(lambda row: highlight_vl_summary(row, bm_ms), axis=1).format(precision=2), 
                             width="stretch", hide_index=True)

            # --- EXPANDER 3: VL MoM Deltas (Expanded Absolute Monthly Data) ---
            with st.expander("📊 VL MoM Performance (Deltas)", expanded=False):
                all_mths = [m["month"] for m in client_data["monthly"]]
                mom_rows = []
                
                for _, row in df_vl.iterrows():
                    vln = row["vl"]
                    vm = vl_monthly.get(vln, {})
                    mom_rec = {
                        "VL Name": vln,
                        "ZM": row.get("ZM", "Unknown"),
                        "Region": row.get("Region", "Unknown")
                    }
                    
                    for mth in all_mths:
                        mom_rec[f"FODs {mth[:3]}"] = vm.get(mth, {}).get("fods", 0) if vm.get(mth) else 0
                    if len(all_mths) >= 2:
                        m1, m2 = all_mths[-2], all_mths[-1]
                        f1 = vm.get(m1, {}).get("fods", 0) if vm.get(m1) else 0
                        f2 = vm.get(m2, {}).get("fods", 0) if vm.get(m2) else 0
                        mom_rec["FOD Growth %"] = round((f2 - f1) / max(f1, 1) * 100, 2) if f1 > 0 else None
                    
                    for ms in ms_list:
                        for mth in all_mths:
                            mom_rec[f"F{ms}% {mth[:3]}"] = vm.get(mth, {}).get(f"pct_{ms}") if vm.get(mth) else None
                        if len(all_mths) >= 2:
                            mom_rec[f"Δ F{ms} (pp)"] = row.get(f"delta_{ms}")
                            
                    if len(all_mths) >= 2:
                        m1, m2 = all_mths[-2], all_mths[-1]
                        mom_rec[f"Median LT {m1[:3]}"] = vm.get(m1, {}).get("median_lt") if vm.get(m1) else None
                        mom_rec[f"Median LT {m2[:3]}"] = vm.get(m2, {}).get("median_lt") if vm.get(m2) else None
                        mom_rec[f"200+% {m1[:3]}"] = vm.get(m1, {}).get("pct_200plus") if vm.get(m1) else None
                        mom_rec[f"200+% {m2[:3]}"] = vm.get(m2, {}).get("pct_200plus") if vm.get(m2) else None
                    
                    mom_rows.append(mom_rec)
                
                if not mom_rows:
                    st.info("No VL MoM data available.")
                else:
                    df_mom = pd.DataFrame(mom_rows)
                    st.dataframe(df_mom.style.apply(highlight_deltas, axis=1).format(precision=2), 
                                 width="stretch", hide_index=True)

            # --- EXPANDER 4: Quality Decline View ---
            with st.expander("📉 VL Quality Decline View", expanded=False):
                all_months = sorted(df_raw["_month"].unique(), key=lambda x: pd.to_datetime("01 " + x))
                if len(all_months) < 2:
                    st.info("Insufficient Month-over-Month data to generate quality decline view.")
                else:
                    curr_m, prev_m = all_months[-1], all_months[-2]
                    ms1, ms2 = CLIENT_DECLINE_MS.get(client, (ms_list[0], key_ms))
                    decline_rows = []
                    
                    for vln in filtered_vl_names:
                        vl_rec = next((r for r in client_data["vl_summary"] if r["vl"] == vln), {})
                        
                        zm = vl_rec.get("ZM", "Unknown")
                        reg = vl_rec.get("Region", "Unknown")
                        cm = vl_rec.get("CM", "Unknown")
                        cl = vl_rec.get("CL", "Unknown")
                        
                        vm = vl_monthly.get(vln, {})
                        curr_d = vm.get(curr_m) or {}
                        prev_d = vm.get(prev_m) or {}
                        
                        curr_fod = curr_d.get("fods")
                        prev_fod = prev_d.get("fods")
                        curr_f1 = curr_d.get(f"pct_{ms1}")
                        prev_f1 = prev_d.get(f"pct_{ms1}")
                        curr_f2 = curr_d.get(f"pct_{ms2}")
                        prev_f2 = prev_d.get(f"pct_{ms2}")
                        
                        d_f1 = round(curr_f1 - prev_f1, 2) if curr_f1 is not None and prev_f1 is not None else None
                        d_f2 = round(curr_f2 - prev_f2, 2) if curr_f2 is not None and prev_f2 is not None else None
                        
                        if curr_fod is not None or prev_fod is not None:
                            decline_rows.append({
                                "VL Name": vln,
                                "ZM Name": zm,
                                "Region": reg,
                                "CM": cm,
                                "CL": cl,
                                f"{curr_m[:3]} MTD FOD": curr_fod if curr_fod is not None else 0,
                                f"LMTD FOD": prev_fod if prev_fod is not None else 0,
                                f"{curr_m[:3]} F{ms1}%": curr_f1,
                                f"LMTD F{ms1}%": prev_f1,
                                f"{curr_m[:3]} F{ms2}%": curr_f2,
                                f"LMTD F{ms2}%": prev_f2,
                                f"Delta F{ms1}": d_f1,
                                f"Delta F{ms2}": d_f2
                            })
                    
                    if decline_rows:
                        df_decline = pd.DataFrame(decline_rows)
                        df_decline = df_decline.sort_values(by=f"Delta F{ms2}", ascending=True, na_position="last")
                        st.dataframe(df_decline.style.apply(highlight_deltas, axis=1).format(precision=2), 
                                     width="stretch", hide_index=True)
                    else:
                        st.info("No records matched quality decline thresholds.")

            # --- EXPANDER 5: Misuse & Anomaly Flags (Last 30 Days) ---
            with st.expander("🚨 VL Misuse & Anomaly Flags (Last 30 Days)", expanded=False):
                # Calculate precise 30-day window metrics
                end_dt = pd.to_datetime(END_DATE)
                start_dt = end_dt - pd.Timedelta(days=30)
                
                df_30d = df_raw[(df_raw["company_name"].str.lower() == client) & 
                                (df_raw["_fod"] >= start_dt) & 
                                (df_raw["_fod"] <= end_dt)]
                
                desired_ms = MISUSE_SHOW_MS.get(client, [key_ms])
                show_ms = [m2 for m2 in desired_ms if m2 in ms_list]
                if key_ms not in show_ms:
                    show_ms = [key_ms] + show_ms
                show_ms = list(dict.fromkeys(show_ms))

                bm_30d = {}
                for m2 in show_ms:
                    col_has = f"has_{m2}"
                    if col_has in df_30d.columns and len(df_30d) > 0:
                        bm_30d[m2] = df_30d[col_has].mean() * 100
                    else:
                        bm_30d[m2] = 0
                bm_med_lt_30d = df_30d["candidate_lifetime_orders_trips"].astype(float).median() if len(df_30d) > 0 else 0

                misuse_rows = []
                for vln in filtered_vl_names:
                    grp = df_30d[df_30d["_vl"] == vln]
                    total_fods = len(grp)
                    if total_fods <= MIN_CURRENT_MTD_FODS: continue

                    zm = grp["ZM"].mode()[0] if not grp["ZM"].empty else "Unknown"
                    reg = grp["Region"].mode()[0] if not grp["Region"].empty else "Unknown"
                    cm = grp["CM"].mode()[0] if not grp["CM"].empty else "Unknown"
                    cl = grp["CL"].mode()[0] if not grp["CL"].empty else "Unknown"

                    lt_all = grp["candidate_lifetime_orders_trips"].astype(float)
                    med_lt = lt_all.median()
                    bel20 = (lt_all < 20).mean() * 100

                    reasons = []
                    sev_scores = []
                    bm_drop_flags = []
                    is_critical_drop = False

                    for m2 in show_ms:
                        col_has = f"has_{m2}"
                        vl_pct = grp[col_has].mean() * 100 if col_has in grp.columns else 0
                        bv = bm_30d.get(m2, 0)
                        if bv > 0:
                            drop_pct = (bv - vl_pct) / bv
                            if drop_pct >= 0.50:  # Drop is >= 50%
                                is_critical_drop = True
                                bm_drop_flags.append(f"F{m2}={vl_pct:.1f}% (≥50% drop from base {bv:.1f}%)")
                            elif drop_pct >= 0.15: # Standard drop
                                bm_drop_flags.append(f"F{m2}={vl_pct:.1f}% (>{15}% drop from base {bv:.1f}%)")
                                sev_scores.append("high")

                    if med_lt < LT_CRITICAL:
                        reasons.append(f"Median LT = {med_lt:.1f} — ghost risk")
                        sev_scores.append("critical")
                    elif med_lt < LT_HIGH:
                        reasons.append(f"Median LT = {med_lt:.1f} — low")
                        sev_scores.append("high")

                    if bel20 > BELOW20_WATCH:
                        reasons.append(f"{bel20:.1f}% <20 LT")
                        sev_scores.append("watch")

                    if not reasons and not bm_drop_flags:
                        continue

                    if is_critical_drop:
                        sev_scores.append("critical")
                    elif not sev_scores and bm_drop_flags:
                        sev_scores.append("watch")

                    final_sev = min(sev_scores, key=lambda s: {"critical": 0, "high": 1, "watch": 2}[s]) if sev_scores else "watch"
                    sev_label = {"critical": "❌ CRITICAL", "high": "🟠 HIGH", "watch": "🟡 WATCH"}[final_sev]

                    combined_reasons = " | ".join(reasons)
                    if bm_drop_flags:
                        combined_reasons += (" | Base Drops: " if combined_reasons else "Base Drops: ") + ", ".join(bm_drop_flags)

                    row_data = {
                        "VL Name": vln,
                        "ZM": zm,
                        "Region": reg,
                        "CM": cm,
                        "CL": cl,
                        "Severity": sev_label,
                        "Total FODs": total_fods,
                        "Median LT": med_lt,
                    }

                    for m2 in show_ms:
                        col_has = f"has_{m2}"
                        vl_pct = grp[col_has].mean() * 100 if col_has in grp.columns else 0
                        bv = bm_30d.get(m2, 0)
                        dropped = (bv > 0) and (vl_pct < bv * 0.85)
                        row_data[f"F{m2}%"] = vl_pct
                        row_data[f"F{m2} Status"] = "⚠️ DROP" if dropped else "✓ OK"

                    row_data["Red Flags"] = combined_reasons
                    misuse_rows.append(row_data)

                if misuse_rows:
                    df_misuse = pd.DataFrame(misuse_rows)
                    severity_map = {"❌ CRITICAL": 0, "🟠 HIGH": 1, "🟡 WATCH": 2}
                    df_misuse["_sev_sort"] = df_misuse["Severity"].map(severity_map)
                    df_misuse = df_misuse.sort_values(by=["_sev_sort", "Total FODs"], ascending=[True, False]).drop(columns=["_sev_sort"])

                    bm_misuse = {
                        "VL Name": "⬛ BENCHMARK (Last 30 Days)",
                        "ZM": "", "Region": "", "CM": "", "CL": "", "Severity": "-",
                        "Total FODs": len(df_30d),
                        "Median LT": bm_med_lt_30d,
                        "Red Flags": "Overall Client Baseline (Last 30 Days)"
                    }
                    for m2 in show_ms:
                        bm_misuse[f"F{m2}%"] = bm_30d.get(m2, 0)
                        bm_misuse[f"F{m2} Status"] = ""

                    df_misuse = pd.concat([pd.DataFrame([bm_misuse]), df_misuse], ignore_index=True)

                    status_cols = [c for c in df_misuse.columns if str(c).endswith("Status")]
                    df_view = df_misuse.drop(columns=status_cols)

                    st.dataframe(df_view.style.apply(highlight_severity_rows, axis=1)
                                                .map(highlight_misuse_status)
                                                .format(precision=2), 
                                 width="stretch", hide_index=True)
                else:
                    st.success("🎉 No vendor anomalies or quality warnings detected for the last 30 days!")

    # --- COMMERCIALS TAB ---
    with tabs[-1]:
        st.header("Financial Revenue Models")
        st.markdown("Automated comparison of Rate Card 1 and Rate Card 2 (Offer 1) vs Prior Month baselines.")
        
        if not fin_data:
            st.info("Insufficient multi-month data to calculate revenue baselines.")
        else:
            for ck, df_fin in fin_data.items():
                with st.expander(f"💳 {ck.title()} Revenue & Share Metrics", expanded=True):
                    
                    colA, colB = st.columns(2)
                    c_regs = ["All"] + sorted([str(x) for x in df_fin["Region"].unique()])
                    sel_c_reg = colA.selectbox(f"Filter Region ({ck.title()})", c_regs, key=f"com_reg_{ck}")
                    
                    df_fin_disp = df_fin.copy()
                    if sel_c_reg != "All":
                        df_fin_disp = df_fin_disp[df_fin_disp["Region"] == sel_c_reg]
                        
                    c_zms = ["All"] + sorted([str(x) for x in df_fin_disp["ZM"].unique()])
                    sel_c_zm = colB.selectbox(f"Filter ZM ({ck.title()})", c_zms, key=f"com_zm_{ck}")
                    
                    if sel_c_zm != "All":
                        df_fin_disp = df_fin_disp[df_fin_disp["ZM"] == sel_c_zm]
                    
                    st.dataframe(df_fin_disp.style.map(style_financials).format(precision=2), width="stretch", hide_index=True)

    # --- SIDEBAR: EXECUTIVE INSIGHTS & ADMIN ---
    with st.sidebar:
        st.header("🤖 AI Insights")
        if not HAS_GEMINI:
            st.warning(
                "AI Module Disabled:\n"
                "Please add `google-genai` to your `requirements.txt`."
            )
        if st.button("Generate Executive Summary"):
            with st.spinner("Analyzing cross-client trends..."):
                summary = draft_summary(results)
                st.markdown(summary)
                
        st.divider()
        
        st.header("🔒 Admin Portal")
        st.caption("Access restricted to authorized personnel.")
        admin_pass = st.text_input("Passkey", type="password")
        
        if admin_pass == "TnQAdmin":
            st.success("Admin access granted.")
            if not HAS_DOCX:
                st.warning("Please add `python-docx` to `requirements.txt` to enable Word Generation.")
            else:
                if st.button("Generate ZM Email Drafts"):
                    with st.spinner("Compiling Word Documents..."):
                        zip_buffer = generate_zm_email_drafts(results, df_raw, mtd_day, END_DATE)
                        st.download_button(
                            label="📥 Download ZM Drafts (.zip)",
                            data=zip_buffer,
                            file_name=f"ZM_Email_Drafts_{END_DATE}.zip",
                            mime="application/zip"
                        )

if __name__ == "__main__":
    main()
