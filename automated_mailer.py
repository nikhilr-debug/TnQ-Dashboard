import os
import smtplib
import zipfile
import io
import time
import requests
import pandas as pd
from datetime import date, timedelta, datetime, timezone
from email.message import EmailMessage
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ==========================================
# 1. CONFIGURATION & TEST MODE
# ==========================================
SENDER_EMAIL = "nikhil.r@vahan.co"
EMAIL_PASSWORD = os.environ.get("EMAIL_APP_PASS")

# TOGGLE THIS: Set to True to route all emails to yourself. Set to False for production.
TEST_MODE = True
TEST_EMAIL = "nikhil.r@vahan.co"

ZM_EMAILS = {
    "Piyush": "piyush.monga@vahan.co",
    "Rohit": "rohit@vahan.co",
    "Vishal": "vishalmittra@vahan.co",
    "Anil Kumar Singh": "anil@vahan.co"
}

# --- REDASH & DATA CONSTANTS ---
REDASH_URL = "https://redash.vahan.link"
QUERY_ID = 17682
REDASH_API_KEY = "4aFm2iOoyx8I91svQccdeZr0jmaiUsMFSRinZcmu"
ACTIVE_CLIENTS = ["blinkit", "swiggy", "swiggy instamart", "uber"]
CLIENT_FULL = {ck: ck.upper() for ck in ACTIVE_CLIENTS}

CLIENT_MS = {
    "blinkit": ["20th", "60th", "100th", "120th", "150th", "200th"],
    "swiggy": ["5th", "10th", "20th", "50th", "60th", "80th", "100th", "150th", "200th"],
    "swiggy instamart": ["5th", "10th", "20th", "50th", "60th", "80th", "100th", "150th", "200th"],
    "uber": ["10th", "20th", "30th", "50th", "100th", "150th", "200th"],
}

CLIENT_KEY_MS = {
    "blinkit": "60th", "swiggy": "20th", "swiggy instamart": "20th", "uber": "20th",
}

CLIENT_DECLINE_MS = {
    "blinkit":          ("20th", "60th"),
    "swiggy":           ("20th", "50th"),
    "swiggy instamart": ("20th", "50th"),
    "uber":             ("10th", "20th"),
}

MIN_VL_FODS = 0
MIN_CURRENT_MTD_FODS = 25
LT_CRITICAL = 5

# Date Calculations
yesterday = date.today() - timedelta(days=1)
mtd_day = yesterday.day
END_DATE = str(yesterday)
IST = timezone(timedelta(hours=5, minutes=30))

# ==========================================
# 2. EMAIL SENDER LOGIC (HTML SUPPORTED)
# ==========================================
def send_email(zm_name, attachment_path, html_body):
    
    # Format Subject: "Quality Report | [Month] MTD [01]-[DD] | [Name]"
    month_name = yesterday.strftime('%B')
    end_day_str = yesterday.strftime('%d')
    subject = f"Quality Report | {month_name} MTD 01-{end_day_str} | {zm_name}"

    if TEST_MODE:
        recipient = TEST_EMAIL
        subject = f"[TEST] {subject}"
        print(f"TEST MODE ACTIVE: Diverting {zm_name}'s email to {recipient}")
    else:
        recipient = ZM_EMAILS.get(zm_name)
        
    if not recipient:
        return

    msg = EmailMessage()
    msg['Subject'] = subject
    msg['From'] = SENDER_EMAIL
    msg['To'] = recipient
    
    # Set the content as HTML so tables render in Gmail/Outlook
    msg.set_content(html_body, subtype='html')

    # Attach the backup Word Document
    if os.path.exists(attachment_path):
        with open(attachment_path, 'rb') as f:
            msg.add_attachment(
                f.read(),
                maintype='application',
                subtype='vnd.openxmlformats-officedocument.wordprocessingml.document',
                filename=os.path.basename(attachment_path)
            )

    try:
        with smtplib.SMTP_SSL('smtp.gmail.com', 465) as smtp:
            smtp.login(SENDER_EMAIL, EMAIL_PASSWORD)
            smtp.send_message(msg)
        print(f"Success: Email sent for ZM {zm_name}")
    except Exception as e:
        print(f"Critical Error sending email for {zm_name}: {e}")

# ==========================================
# 3. DATA FETCHING & PROCESSING (Unchanged)
# ==========================================
def get_daily_refresh_key():
    now = datetime.now(IST)
    if now.hour < 13 or (now.hour == 13 and now.minute < 30):
        return str(now.date() - timedelta(days=1))
    return str(now.date())

def fetch_redash(refresh_key):
    body_fresh = {"parameters": {"Client": ACTIVE_CLIENTS}, "max_age": 7200}
    r = requests.post(f"{REDASH_URL}/api/queries/{QUERY_ID}/results?api_key={REDASH_API_KEY}", json=body_fresh, timeout=30)
    j = r.json()
    if "query_result" in j: return j["query_result"]["data"]["rows"]
    
    body_cached = {**body_fresh, "max_age": 7200}
    for _ in range(40):
        time.sleep(15)
        r2 = requests.post(f"{REDASH_URL}/api/queries/{QUERY_ID}/results?api_key={REDASH_API_KEY}", json=body_cached, timeout=30)
        j2 = r2.json()
        if "query_result" in j2: return j2["query_result"]["data"]["rows"]
    return []

def run_analysis(rows):
    if not rows: return {}, pd.DataFrame()
    df = pd.DataFrame(rows)
    df["_fod"] = pd.to_datetime(df["first_date_of_work"], format="%Y-%m-%d", errors="coerce")
    valid = df["_fod"].notna() & (df["_fod"].dt.day <= mtd_day) & (df["_fod"] <= pd.Timestamp(END_DATE))
    df = df[valid].copy()
    df["_month"] = df["_fod"].dt.strftime("%b-%Y")
    df = df.drop_duplicates(subset=["phone_number", "_month"])
    df["_vl"] = df["vl_name"].fillna("Unknown")
    col_map = {str(c).strip().lower(): c for c in df.columns}
    df["ZM"] = df[col_map["zm"]].fillna("Unknown") if "zm" in col_map else "Unknown"

    results = {}
    for client in ACTIVE_CLIENTS:
        sub = df[df["company_name"].str.lower() == client].copy()
        ms_list = CLIENT_MS.get(client, [])
        key_ms = CLIENT_KEY_MS.get(client, ms_list[0])
        for ms in ms_list:
            col = f"{ms}_order_date"
            if col not in sub.columns: sub[col] = None
            sub[col + "_dt"] = pd.to_datetime(sub[col], format="%Y-%m-%d", errors="coerce")
            sub[f"has_{ms}"] = ((sub[col + "_dt"].dt.year == sub["_fod"].dt.year) & (sub[col + "_dt"].dt.month == sub["_fod"].dt.month) & (sub[col + "_dt"].dt.day <= mtd_day)).astype(int)

        all_months = sorted(sub["_month"].unique(), key=lambda x: pd.to_datetime("01 " + x))
        monthly = []
        for m in all_months:
            g = sub[sub["_month"] == m]
            if len(g) == 0: continue
            rec = {"month": m, "fods": len(g)}
            for ms in ms_list: rec[f"pct_{ms}"] = round(g[f"has_{ms}"].mean() * 100, 2)
            monthly.append(rec)
        bm_ms = {ms2: round(sum(m.get(f"pct_{ms2}", 0) for m in monthly) / max(len(monthly), 1), 2) for ms2 in ms_list}

        vl_summary = []; vl_monthly = {}
        for vl_name, vl_df in sub.groupby("_vl"):
            if len(vl_df) < MIN_VL_FODS: continue
            lt_all = vl_df["candidate_lifetime_orders_trips"].astype(float)
            rec = {"vl": vl_name, "ZM": vl_df["ZM"].mode()[0] if not vl_df["ZM"].empty else "Unknown", "total_fods": len(vl_df), "median_lt": round(lt_all.median(), 2)}
            for ms in ms_list: rec[f"pct_{ms}"] = round(vl_df[f"has_{ms}"].mean() * 100, 2)
            vm = {}
            for m in all_months:
                m_df = vl_df[vl_df["_month"] == m]
                if len(m_df) < 5: continue
                m_rec = {"fods": len(m_df)}
                for ms in ms_list: m_rec[f"pct_{ms}"] = round(m_df[f"has_{ms}"].mean() * 100, 2)
                vm[m] = m_rec
            vl_monthly[vl_name] = vm
            curr_m = all_months[-1] if all_months else None
            rec["curr_m_fods"] = vm[curr_m]["fods"] if curr_m and vm.get(curr_m) else 0
            vl_summary.append(rec)

        vl_summary = sorted(vl_summary, key=lambda x: x.get("curr_m_fods", 0), reverse=True)
        results[client] = {"monthly": monthly, "vl_summary": vl_summary, "vl_monthly": vl_monthly, "bm_ms": bm_ms, "milestones": ms_list, "key_ms": key_ms}
    return results

# ==========================================
# 4. WORD DOC & HTML GENERATION
# ==========================================
def set_thick_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = tblPr.first_child_found_in("w:tblBorders")
    if tblBorders is None:
        tblBorders = OxmlElement('w:tblBorders')
        tblPr.append(tblBorders)
    else: tblBorders.clear()
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single'); border.set(qn('w:sz'), '12'); border.set(qn('w:space'), '0'); border.set(qn('w:color'), '000000')
        tblBorders.append(border)

def _fmt_pct_word(val): return "-" if pd.isna(val) or val is None else f"{val:.1f}%"

def generate_docs_and_html(results):
    output_dir = "temp_zm_drafts"
    os.makedirs(output_dir, exist_ok=True)
    html_payloads = {}
    
    unique_zms = set()
    for ck in ACTIVE_CLIENTS:
        if ck in results:
            for vl in results[ck]["vl_summary"]:
                if vl.get("ZM") and vl["ZM"] != "Unknown" and vl["ZM"] in ZM_EMAILS.keys():
                    unique_zms.add(str(vl["ZM"]).strip())
    
    cohort_month = yesterday.strftime('%B')
    
    for zm_name in unique_zms:
        doc = Document()
        for section in doc.sections:
            section.top_margin = Pt(36); section.bottom_margin = Pt(36); section.left_margin = Pt(36); section.right_margin = Pt(36)
            
        doc.add_paragraph(f"Hi {zm_name},").runs[0].bold = True
        doc.add_paragraph(f"Please find {cohort_month}'s TnQ quality Report for your cluster at the client level below.")

        # --- HTML TEMPLATE SETUP ---
        html_body = f"""
        <html>
        <head>
        <style>
            body {{ font-family: Arial, sans-serif; font-size: 14px; color: #333; }}
            table {{ border-collapse: collapse; width: 100%; margin-bottom: 20px; font-size: 12px; }}
            th, td {{ border: 1px solid #000; padding: 6px 8px; text-align: left; }}
            th {{ background-color: #f2f2f2; font-weight: bold; }}
            .right-align {{ text-align: right; }}
            h2 {{ color: #C55A00; margin-top: 20px; margin-bottom: 10px; font-size: 18px; }}
            h3 {{ color: #333; margin-top: 15px; margin-bottom: 5px; font-size: 14px; text-decoration: underline; }}
        </style>
        </head>
        <body>
            <p><strong>Hi {zm_name},</strong></p>
            <p>Please find {cohort_month}'s TnQ quality Report for your cluster at the client level below.</p>
        """
        
        for ck in ACTIVE_CLIENTS:
            if ck not in results: continue
            data = results[ck]
            client_label = CLIENT_FULL.get(ck, ck.upper())
            mon = data["monthly"]
            if len(mon) < 2: continue
            curr_m, prev_m = mon[-1]["month"], mon[-2]["month"]
            ms1, ms2 = CLIENT_DECLINE_MS.get(ck, (data["milestones"][0], data["key_ms"]))

            t1_rows = []
            for vl in data["vl_summary"]:
                if vl.get("ZM") != zm_name: continue
                vln = vl["vl"]
                vm = data["vl_monthly"].get(vln, {})
                curr_d, prev_d = vm.get(curr_m) or {}, vm.get(prev_m) or {}
                curr_f1, prev_f1 = curr_d.get(f"pct_{ms1}"), prev_d.get(f"pct_{ms1}")
                curr_f2, prev_f2 = curr_d.get(f"pct_{ms2}"), prev_d.get(f"pct_{ms2}")
                d_f1 = round(curr_f1 - prev_f1, 1) if curr_f1 is not None and prev_f1 is not None else None
                d_f2 = round(curr_f2 - prev_f2, 1) if curr_f2 is not None and prev_f2 is not None else None

                if (d_f1 is not None and d_f1 < 0) or (d_f2 is not None and d_f2 < 0):
                    t1_rows.append([str(vln), str(zm_name), f"{curr_d.get('fods', 0):,}", f"{prev_d.get('fods', 0):,}", _fmt_pct_word(curr_f1), _fmt_pct_word(prev_f1), _fmt_pct_word(curr_f2), _fmt_pct_word(prev_f2), f"{d_f1:+.1f}%" if d_f1 is not None else "-", f"{d_f2:+.1f}%" if d_f2 is not None else "-"])

            t2_ms_list = ["20th", "60th", "100th", "200th"]
            t2_rows = []
            for vl_rec in data["vl_summary"]:
                if vl_rec.get("ZM") != zm_name: continue
                total_fods = vl_rec.get("total_fods", 0)
                if total_fods <= MIN_CURRENT_MTD_FODS: continue
                
                med_lt = vl_rec.get("median_lt", 999)
                is_critical = False
                red_flags = []
                
                if med_lt < LT_CRITICAL:
                    red_flags.append(f"Median LT = {med_lt:.1f} risk"); is_critical = True
                
                for m2 in t2_ms_list:
                    if m2 in data["milestones"]:
                        vl_pct = vl_rec.get(f"pct_{m2}", 0)
                        bv = data["bm_ms"].get(m2, 0)
                        if bv > 0:
                            drop_pct = (bv - vl_pct) / bv
                            if drop_pct >= 0.50: red_flags.insert(0, f"Critical Drop F{m2}={vl_pct:.1f}%"); is_critical = True
                            elif drop_pct >= 0.15: red_flags.append(f"Drop F{m2}={vl_pct:.1f}%")
                                
                if is_critical:
                    row_data = [str(vl_rec['vl']), str(zm_name), "CRITICAL", f"{total_fods:,}", str(round(med_lt, 1))]
                    for m2 in t2_ms_list:
                        if m2 in data["milestones"]:
                            row_data.append(_fmt_pct_word(vl_rec.get(f"pct_{m2}", 0)))
                            row_data.append(_fmt_pct_word(data["bm_ms"].get(m2, 0)))
                        else: row_data.extend(["-", "-"])
                    row_data.append(" | ".join(red_flags))
                    t2_rows.append(row_data)

            if t1_rows or t2_rows:
                client_h = doc.add_heading(client_label, level=2)
                client_h.runs[0].font.color.rgb = RGBColor(197, 90, 0)
                html_body += f"<h2>{client_label}</h2>"

                if t1_rows:
                    doc.add_heading("MTD VS LMD report", level=3)
                    html_body += f"<h3>MTD VS LMD report</h3><table><tr>"
                    t1_headers = ["VL Name", "ZM Name", f"{curr_m[:3]} MTD", f"LMTD FOD", f"MTD F{ms1}%", f"LMTD F{ms1}%", f"MTD F{ms2}%", f"LMTD F{ms2}%", f"Delta F{ms1}", f"Delta F{ms2}"]
                    
                    table = doc.add_table(rows=1, cols=len(t1_headers))
                    set_thick_borders(table)
                    for i, h in enumerate(t1_headers): 
                        table.rows[0].cells[i].text = h
                        html_body += f"<th>{h}</th>"
                    html_body += "</tr>"
                    
                    for row_data in t1_rows:
                        row_cells = table.add_row().cells
                        html_body += "<tr>"
                        for i, val in enumerate(row_data): 
                            row_cells[i].text = str(val)
                            # Align numbers to right in HTML
                            css_class = ' class="right-align"' if i >= 2 else ''
                            html_body += f"<td{css_class}>{val}</td>"
                        html_body += "</tr>"
                    doc.add_paragraph()
                    html_body += "</table>"

                if t2_rows:
                    doc.add_heading("Platform Avg vs VL Performance (MTD)", level=3)
                    html_body += f"<h3>Platform Avg vs VL Performance (MTD)</h3><table><tr>"
                    t2_headers = ["VL Name", "ZM", "Severity", "Total FODs", "Median LT", "F20th%\n(Overall)", "F20th%\n(Base)", "F60th%\n(Overall)", "F60th%\n(Base)", "F100th%\n(Overall)", "F100th%\n(Base)", "F200th%\n(Overall)", "F200th%\n(Base)", "Red Flags"]
                    
                    table = doc.add_table(rows=1, cols=len(t2_headers))
                    set_thick_borders(table)
                    for i, h in enumerate(t2_headers): 
                        table.rows[0].cells[i].text = h
                        html_body += f"<th>{h.replace(chr(10), '<br>')}</th>"
                    html_body += "</tr>"

                    for row_data in t2_rows:
                        row_cells = table.add_row().cells
                        html_body += "<tr>"
                        for i, val in enumerate(row_data): 
                            row_cells[i].text = str(val)
                            css_class = ' class="right-align"' if 3 <= i <= 12 else ''
                            html_body += f"<td{css_class}>{val}</td>"
                        html_body += "</tr>"
                    doc.add_paragraph()
                    html_body += "</table>"

        html_body += "</body></html>"
        html_payloads[zm_name] = html_body

        safe_zm_name = "".join([c for c in zm_name if c.isalpha() or c.isdigit() or c==' ']).rstrip().replace(' ', '_')
        file_path = os.path.join(output_dir, f"ZM_Report_{safe_zm_name}.docx")
        doc.save(file_path)
        
    return output_dir, html_payloads

# ==========================================
# 5. MAIN EXECUTION
# ==========================================
def run_automation():
    print("Starting Automated Mailer Job...")
    refresh_key = get_daily_refresh_key()
    rows = fetch_redash(refresh_key)
    if not rows:
        print("Aborting: No data returned from Redash.")
        return
        
    results = run_analysis(rows)
    output_dir, html_payloads = generate_docs_and_html(results)
    
    for filename in os.listdir(output_dir):
        if not filename.endswith(".docx"): continue
        file_path = os.path.join(output_dir, filename)
        
        target_zm = None
        for zm_key in ZM_EMAILS.keys():
            safe_key = "".join([c for c in zm_key if c.isalpha() or c.isdigit() or c==' ']).rstrip().replace(' ', '_')
            if safe_key in filename:
                target_zm = zm_key
                break
                
        if target_zm:
            # Fetch the pre-rendered HTML body for this specific ZM
            email_body_html = html_payloads.get(target_zm, "<html><body><p>Error generating report content.</p></body></html>")
            send_email(zm_name=target_zm, attachment_path=file_path, html_body=email_body_html)
            
if __name__ == "__main__":
    run_automation()
